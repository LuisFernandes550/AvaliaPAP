from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import AreaPAP

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

ROTULOS_ALUNO = (
    r"aluno(?:\(a\))?",
    r"nome\s+aluno",
    r"nome(?:\s+do\s+aluno)?",
    r"estudante",
    r"autor(?:\(a\))?(?:es)?",
    r"candidato(?:\(a\))?",
    r"elaborado\s+por",
)

ROTULOS_TITULO = (
    r"t[ií]tulo(?:\s+do\s+projeto)?",
    r"designa[cç][aã]o(?:\s+do\s+projeto)?",
    r"nome\s+do\s+projeto",
    r"projeto",
    r"tema(?:\s+do\s+projeto)?",
)

MARCADORES_PAP = (
    r"prova\s+de\s+aptid[aã]o\s+profissional",
    r"prova\s+de\s+aptid[aã]o\s+profissional\s*\(?\s*pap\s*\)?",
    r"relat[oó]rio\s+(?:da\s+)?(?:prova\s+de\s+)?aptid[aã]o",
    r"relat[oó]rio\s+(?:final\s+)?(?:da\s+)?pap",
)

LINHAS_IGNORAR_TITULO = (
    r"^prova\s+de\s+aptid",
    r"^relat[oó]rio",
    r"^pap$",
    r"^escola\b",
    r"^agrupamento",
    r"^curso\b",
    r"^aluno",
    r"^orientador",
    r"^ano\s+letivo",
    r"^disciplina",
    r"^m[oó]dulo",
    r"^turma",
    r"^n[úu]mero",
    r"^n\.?[ºo°]",
    r"^n\.?\s*[ºo°]?\s*:?\s*\d",
    r"^professor",
    r"^data\b",
    r"^local\b",
    r"^índice",
    r"^indice",
    r"^agrade",
    r"^dedicat",
    r"^resumo",
    r"^abstract",
    r"^coordenador",
    r"^diretor",
    r"^formando",
)

PADROES_FICHEIRO_ALUNO = [
    r"(?i)relat[oó]rio\s*pap[_\s-]+(.+)$",
    r"(?i)relat[oó]riopap[_\s-]+(.+)$",
    r"(?i)pap[_\s-]+(.+)$",
    r"(?i)relat[oó]rio[_\s-]+(.+)$",
]

INDICADORES_AREA: dict[AreaPAP, tuple[str, ...]] = {
    AreaPAP.WEBSITE: (
        "website", "site web", "pagina web", "página web", "html", "css",
        "wordpress", "front-end", "frontend", "backend", "web app",
    ),
    AreaPAP.APLICACAO_MOVEL: (
        "aplicação móvel", "aplicacao movel", "app móvel", "app movel",
        "android", "ios", "flutter", "react native", "kotlin", "swift",
        "mobile", "smartphone",
    ),
    AreaPAP.ROBOTICA: (
        "robótica", "robotica", "robot", "arduino", "raspberry", "sensor",
        "atuador", "motores", "iot", "automação", "automacao",
    ),
    AreaPAP.JOGO: (
        "jogo", "game", "unity", "godot", "unreal", "jogabilidade",
        "gameplay", "nível", "nivel", "personagem",
    ),
}


def _normalizar(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def _celulas_unicas(row) -> list[str]:
    vistas: list[str] = []
    for cell in row.cells:
        t = _normalizar(cell.text)
        if t and t not in vistas:
            vistas.append(t)
    return vistas


def _corresponde_rotulo(texto: str, padroes: tuple[str, ...]) -> bool:
    t = texto.lower().strip()
    return any(re.fullmatch(p, t, re.I) or re.match(rf"^{p}\s*:", t, re.I) for p in padroes)


def _limpar_nome(valor: str) -> str:
    nome = _normalizar(valor)
    nome = re.sub(r"\s*[Nn]\.?\s*[ºo°]\s*:?\s*\d+.*$", "", nome, flags=re.I).strip()
    nome = re.sub(
        r"\s*(?:Professor\s+Orientador|Ano\s+Letivo|Curso\s+Profissional).*$",
        "",
        nome,
        flags=re.I,
    ).strip()
    nome = re.sub(r"\s*-\s*$", "", nome).strip()
    return nome


def _titulo_de_linha_pap(linha: str) -> str:
    m = re.search(
        r"prova\s+de\s+aptid[aã]o\s+profissional\s*(.+?)(?:\s*(?:curso|nome\s+aluno)|$)",
        linha,
        re.I,
    )
    if m:
        titulo = _normalizar(m.group(1))
        titulo = re.sub(r"^[\(\)\-–—:\s]+", "", titulo)
        if titulo and not _deve_ignorar_como_titulo(titulo):
            return titulo
    return ""


def _extrair_capa_prioritaria(linhas: list[str]) -> tuple[str, str]:
    nome = ""
    for linha in linhas:
        lower = linha.lower()
        if "nome aluno" in lower or "nome do aluno" in lower:
            nome = nome or _extrair_nome_de_texto(linha)
        if "nome completo" in lower:
            m = re.search(r"nome\s+completo\s*:\s*([^|\n]+)", linha, re.I)
            if m:
                nome = nome or _limpar_nome(m.group(1))

    # Titulo: linha da capa logo a seguir ao marcador PAP (ordem do documento)
    titulo = _extrair_titulo_apos_pap(linhas)
    if not titulo:
        for linha in linhas:
            if "prova de aptid" in linha.lower():
                t = _titulo_de_linha_pap(linha)
                if t:
                    titulo = t
                    break
    return nome, titulo


def _separar_palavras_nome(nome: str) -> str:
    nome = nome.replace("_", " ").replace("-", " ")
    nome = re.sub(r"(?<=[a-záàâãéèêíìîóòôõúùç])(?=[A-ZÁÀÂÃÉÈÊÍÌÎÓÒÔÕÚÙÇ])", " ", nome)
    nome = re.sub(r"(?<=[A-ZÁÀÂÃÉÈÊÍÌÎÓÒÔÕÚÙÇ])(?=[A-ZÁÀÂÃÉÈÊÍÌÎÓÒÔÕÚÙÇ][a-záàâãéèêíìîóòôõúùç])", " ", nome)
    return _normalizar(nome)


def _extrair_nome_de_texto(texto: str) -> str:
    t = _normalizar(texto)
    for padrao in ROTULOS_ALUNO:
        m = re.search(rf"{padrao}\s*:\s*([^|\n]+)", t, re.I)
        if m:
            nome = _limpar_nome(m.group(1))
            if nome and len(nome) > 2:
                return nome
    return ""


def _valor_de_rotulo(texto: str) -> str:
    for sep in (":", " - ", " – ", " — "):
        if sep in texto:
            parte = texto.split(sep, 1)[1].strip()
            if parte and not _corresponde_rotulo(parte, ROTULOS_ALUNO + ROTULOS_TITULO):
                if any(re.search(rf"^{p}\s*:", parte, re.I) for p in ROTULOS_ALUNO):
                    nome = _extrair_nome_de_texto(parte)
                    return nome or parte
                return parte
    return ""


def _extrair_de_tabelas_capa(doc: Document) -> tuple[str, str]:
    nome = ""
    titulo = ""
    for tabela in doc.tables[:8]:
        linhas = []
        for row in tabela.rows:
            celulas = _celulas_unicas(row)
            if celulas:
                linhas.append(celulas)

        for celulas in linhas:
            texto_linha = " | ".join(celulas)
            nome = nome or _extrair_nome_de_texto(texto_linha)
            for padrao in ROTULOS_ALUNO:
                for i, cel in enumerate(celulas):
                    if re.search(rf"^{padrao}\s*:?\s*$", cel, re.I):
                        if i + 1 < len(celulas):
                            nome = nome or celulas[i + 1]
                    val = _valor_de_rotulo(cel)
                    if val and re.search(rf"^{padrao}\s*:", cel, re.I):
                        nome = nome or _limpar_nome(val)
            for padrao in ROTULOS_TITULO:
                for i, cel in enumerate(celulas):
                    if re.search(rf"^{padrao}\s*:?\s*$", cel, re.I):
                        if i + 1 < len(celulas):
                            candidato = celulas[i + 1]
                            if not _deve_ignorar_como_titulo(candidato):
                                titulo = titulo or candidato
                    val = _valor_de_rotulo(cel)
                    if val and re.search(rf"^{padrao}\s*:", cel, re.I):
                        if not _deve_ignorar_como_titulo(val):
                            titulo = titulo or val

        titulo_tab = _extrair_titulo_apos_pap(linhas)
        if titulo_tab:
            titulo = titulo or titulo_tab

    return nome, titulo


def _textos_de_shape(shape) -> list[str]:
    textos: list[str] = []
    try:
        if shape.TextFrame.HasText:
            texto = shape.TextFrame.TextRange.Text.strip()
            texto = texto.replace("\r", "").replace("\x07", "").strip()
            if texto:
                textos.append(texto)
    except Exception:
        pass
    try:
        if shape.GroupItems.Count > 0:
            for i in range(1, shape.GroupItems.Count + 1):
                textos.extend(_textos_de_shape(shape.GroupItems(i)))
    except Exception:
        pass
    return textos


def _paragrafos_do_xml(caminho: Path) -> list[str]:
    """Le todos os paragrafos do .docx via XML, incluindo caixas de texto (shapes).

    Nao depende do Microsoft Word e capta a capa mesmo quando o texto esta
    dentro de caixas de texto que o python-docx nao le.
    """
    partes: list[str] = []
    try:
        with zipfile.ZipFile(caminho) as z:
            nomes_disponiveis = set(z.namelist())
            alvos = ["word/document.xml"]
            alvos += sorted(
                n for n in nomes_disponiveis
                if re.match(r"word/(header|footer)\d*\.xml$", n)
            )
            for nome in alvos:
                if nome not in nomes_disponiveis:
                    continue
                try:
                    root = ET.fromstring(z.read(nome))
                except ET.ParseError:
                    continue

                buffer: list[str] = []

                def _flush() -> None:
                    texto = _normalizar("".join(buffer))
                    buffer.clear()
                    if texto:
                        partes.append(texto)

                for el in root.iter():
                    tag = el.tag
                    if tag == _W_NS + "p":
                        _flush()
                    elif tag == _W_NS + "t":
                        if el.text:
                            buffer.append(el.text)
                    elif tag in (_W_NS + "tab", _W_NS + "br", _W_NS + "cr"):
                        buffer.append(" ")
                _flush()
    except Exception:
        return partes

    vistos: set[str] = set()
    unicos: list[str] = []
    for p in partes:
        if p not in vistos:
            vistos.add(p)
            unicos.append(p)
    return unicos


def _texto_capa_word_com(caminho: Path) -> str:
    if sys.platform != "win32":
        return ""
    word = None
    doc = None
    try:
        import pythoncom
        import win32com.client  # type: ignore

        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(caminho.resolve()), ReadOnly=True, AddToRecentFiles=False)
        partes: list[str] = []

        try:
            conteudo = doc.Content.Text[:8000].replace("\r", "\n").replace("\x07", "")
            for linha in conteudo.split("\n"):
                linha = linha.strip()
                if linha and linha not in partes:
                    partes.append(linha)
        except Exception:
            pass

        if doc.Paragraphs.Count > 0:
            fim = min(100, doc.Paragraphs.Count)
            for i in range(1, fim + 1):
                t = doc.Paragraphs(i).Range.Text.strip()
                if t:
                    partes.append(t.replace("\r", "").replace("\x07", ""))

        for ti in range(1, doc.Tables.Count + 1):
            tabela = doc.Tables(ti)
            for ri in range(1, tabela.Rows.Count + 1):
                celulas: list[str] = []
                for ci in range(1, tabela.Columns.Count + 1):
                    try:
                        texto = tabela.Cell(ri, ci).Range.Text.strip()
                        texto = texto.replace("\r", "").replace("\x07", "").strip()
                        if texto and texto not in celulas:
                            celulas.append(texto)
                    except Exception:
                        continue
                if celulas:
                    partes.append(" | ".join(celulas))

        for si in range(1, doc.Shapes.Count + 1):
            try:
                partes.extend(_textos_de_shape(doc.Shapes(si)))
            except Exception:
                continue

        vistos: set[str] = set()
        unicos: list[str] = []
        for p in partes:
            if p not in vistos:
                vistos.add(p)
                unicos.append(p)
        return "\n".join(unicos)
    except Exception:
        return ""
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            import pythoncom

            pythoncom.CoUninitialize()
        except Exception:
            pass


def extrair_texto_docx(caminho: Path) -> str:
    doc = Document(str(caminho))
    partes: list[str] = []
    for par in doc.paragraphs:
        t = par.text.strip()
        if t:
            partes.append(t)
    for tabela in doc.tables:
        for row in tabela.rows:
            celulas = _celulas_unicas(row)
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def extrair_estrutura_docx(caminho: Path) -> list[str]:
    """Lista os titulos/capitulos do relatorio (paragrafos com estilo de titulo)."""
    doc = Document(str(caminho))
    titulos: list[str] = []
    for par in doc.paragraphs:
        t = par.text.strip()
        if not t:
            continue
        estilo = (par.style.name or "").lower() if par.style else ""
        eh_titulo = "titulo" in estilo or "título" in estilo or "heading" in estilo
        # Fallback: linhas numeradas tipo "3.", "3.1", "3.1.1"
        if not eh_titulo and re.match(r"^\d+(\.\d+){0,2}\.?\s+\S", t):
            eh_titulo = True
        if eh_titulo:
            # Remove numero de pagina do indice (ex.: "3.1. Titulo\t23" ou "Titulo .... 23")
            t = re.sub(r"[\s\.\t]*\d{1,4}$", "", t).strip(" .\t")
            if t and len(t) <= 120:
                titulos.append(t)
    # Remove duplicados preservando ordem
    vistos = set()
    unicos = []
    for t in titulos:
        chave = t.lower()
        if chave not in vistos:
            vistos.add(chave)
            unicos.append(t)
    return unicos


def texto_para_avaliacao(caminho: Path) -> str:
    """Texto do relatorio com a ESTRUTURA (indice) no inicio, para a IA ter sempre
    o mapa de topicos mesmo que o corpo seja truncado."""
    estrutura = extrair_estrutura_docx(caminho)
    corpo = extrair_texto_docx(caminho)
    if not estrutura:
        return corpo
    indice = "\n".join(estrutura)
    return f"ESTRUTURA DO RELATORIO (titulos/capitulos presentes):\n{indice}\n\nCONTEUDO:\n{corpo}"


def _nome_de_ficheiro(nome_ficheiro: str) -> str:
    stem = Path(nome_ficheiro).stem
    for padrao in PADROES_FICHEIRO_ALUNO:
        m = re.search(padrao, stem)
        if m:
            nome = _separar_palavras_nome(m.group(1))
            nome = re.sub(r"(?i)relat[oó]rio|pap", "", nome).strip()
            if nome:
                return _normalizar(nome)
    return ""


def _eh_marcador_pap(texto: str) -> bool:
    lower = texto.lower().strip()
    return any(re.search(p, lower, re.I) for p in MARCADORES_PAP)


def _parece_definicao_pap(lower: str) -> bool:
    return bool(
        re.search(
            r"consiste\s+na\s+apresenta|perante\s+um\s+j[uú]ri"
            r"|apresenta\w*\s+e\s+defesa|\(\s*pap\s*\)",
            lower,
        )
    )


def _deve_ignorar_como_titulo(texto: str) -> bool:
    lower = texto.lower().strip()
    if not lower or len(lower) < 3:
        return True
    if len(lower) > 70:  # titulos de PAP sao curtos (nome do produto/projeto)
        return True
    if _eh_marcador_pap(texto):
        return True
    if _parece_definicao_pap(lower):
        return True
    if _corresponde_rotulo(texto, ROTULOS_ALUNO):
        return True
    if re.search(r"^(?:nome\s+aluno|professor\s+orientador|ano\s+letivo|curso)\s*:", lower, re.I):
        return True
    if re.search(r"agrupamento|escola|tgpsi|curso\s+profissional", lower, re.I) and len(lower) > 25:
        return True
    # Nome do curso (por vezes parte em duas linhas na capa)
    if re.search(
        r"programa\w*\s+de\s+sistemas|sistemas\s+inform[aá]ticos"
        r"|t[eé]cnico\s+de\s+gest|^e\s+programa",
        lower,
    ):
        return True
    if re.search(r"^pdf\b|portable document", lower, re.I):
        return True
    # Linhas do indice (terminam com numero de pagina) ou dedicatoria/agradecimentos
    if re.search(r"\s\d{1,3}$", lower) and re.search(r"^\d|introdu|�ndice|indice|figura|tabela", lower):
        return True
    return any(re.search(p, lower, re.I) for p in LINHAS_IGNORAR_TITULO)


def _extrair_titulo_apos_pap(linhas: list) -> str:
    textos = []
    for item in linhas:
        if isinstance(item, list):
            textos.extend(item)
        else:
            textos.append(str(item))

    for i, texto in enumerate(textos):
        if not _eh_marcador_pap(texto):
            continue
        # So o texto imediatamente a seguir ao marcador PAP da capa (nao definicoes)
        if _parece_definicao_pap(texto.lower()):
            continue
        for j in range(i + 1, min(i + 8, len(textos))):
            linha = _normalizar(str(textos[j]))
            if not linha:
                continue
            for padrao in ROTULOS_TITULO:
                if re.search(rf"^{padrao}\s*:", linha, re.I):
                    val = _valor_de_rotulo(linha)
                    if val and not _deve_ignorar_como_titulo(val):
                        return val
            if re.search(r"^aluno\s*:|^orientador\s*:|^nome\s+aluno|^curso\s*:", linha, re.I):
                break
            if _deve_ignorar_como_titulo(linha):
                continue  # salta artefactos (ex.: "sS"), curso, etc.
            # Primeira linha valida apos o marcador PAP e o titulo
            return _normalizar(linha)
        # So consideramos o primeiro marcador PAP (o da capa)
        break
    return ""


def _extrair_de_paragrafos(paragrafos: list[str]) -> tuple[str, str]:
    nome = ""
    titulo = _extrair_titulo_apos_pap(paragrafos)

    for t in paragrafos[:80]:
        if not nome:
            nome = _extrair_nome_de_texto(t) or nome
            for padrao in ROTULOS_ALUNO:
                if re.search(rf"^{padrao}\s*:", t, re.I):
                    val = _valor_de_rotulo(t)
                    nome = _limpar_nome(val) if val else nome
        if not titulo:
            for padrao in ROTULOS_TITULO:
                if re.search(rf"^{padrao}\s*:", t, re.I):
                    val = _valor_de_rotulo(t)
                    if val and not _deve_ignorar_como_titulo(val):
                        titulo = val

    return nome, titulo


def _detetar_area(texto: str) -> AreaPAP:
    lower = texto.lower()
    scores = {a: 0 for a in INDICADORES_AREA}
    for area, palavras in INDICADORES_AREA.items():
        for p in palavras:
            if p in lower:
                scores[area] += 1
    melhor = max(scores, key=scores.get)
    return melhor if scores[melhor] >= 1 else AreaPAP.NAO_DETETADA


def _juntar_linhas_unicas(*fontes: list[str]) -> list[str]:
    vistas: set[str] = set()
    resultado: list[str] = []
    for fonte in fontes:
        for linha in fonte:
            linha = linha.strip()
            if linha and linha not in vistas:
                vistas.add(linha)
                resultado.append(linha)
    return resultado


def analisar_relatorio(caminho: Path) -> dict:
    """Extrai nome do aluno, titulo da PAP e area a partir da capa do relatorio."""
    caminho = Path(caminho)
    doc = Document(str(caminho))
    texto_completo = extrair_texto_docx(caminho)

    nome_tab, titulo_tab = _extrair_de_tabelas_capa(doc)

    # Fonte principal: XML do .docx (le caixas de texto/shapes da capa)
    paragrafos_xml = _paragrafos_do_xml(caminho)[:150]
    paragrafos_docx = [p.text.strip() for p in doc.paragraphs[:80] if p.text.strip()]
    linhas_texto = [l.strip() for l in texto_completo.split("\n")[:120] if l.strip()]

    paragrafos = _juntar_linhas_unicas(paragrafos_xml, paragrafos_docx, linhas_texto)
    nome_capa, titulo_capa = _extrair_capa_prioritaria(paragrafos)
    nome_par, titulo_par = _extrair_de_paragrafos(paragrafos)

    nome = nome_capa or nome_tab or nome_par
    titulo = titulo_capa or titulo_tab or titulo_par

    # Ultimo recurso para o NOME: Word COM (o titulo vem apenas da capa/XML)
    if not nome and sys.platform == "win32":
        texto_com = _texto_capa_word_com(caminho)
        if texto_com:
            for linha in texto_com.split("\n"):
                nome = _extrair_nome_de_texto(linha)
                if nome:
                    break

    if not nome:
        for linha in linhas_texto:
            nome = _extrair_nome_de_texto(linha)
            if nome:
                break

    nome_ficheiro = _nome_de_ficheiro(caminho.name)
    if not nome:
        nome = nome_ficheiro
    elif nome_ficheiro and nome.replace(" ", "").lower() == nome_ficheiro.replace(" ", "").lower():
        # Preferir nome do ficheiro se o extraido for igual sem espacos (ex: GabrielBispo)
        nome = nome_ficheiro

    if not nome:
        nome = _separar_palavras_nome(caminho.stem)
    if not titulo:
        titulo = "Titulo nao identificado - confirmar manualmente"
    elif titulo == titulo.lower():
        # Titulo todo em minusculas -> capitaliza (ex.: "sentinela vermelho")
        titulo = titulo.title()

    area = _detetar_area(texto_completo[:12000])

    estrutura = extrair_estrutura_docx(caminho)
    if estrutura:
        indice = "\n".join(estrutura)
        texto_para_ia = (
            f"ESTRUTURA DO RELATORIO (titulos/capitulos presentes):\n{indice}\n\n"
            f"CONTEUDO:\n{texto_completo}"
        )
    else:
        texto_para_ia = texto_completo

    return {
        "nome": _normalizar(nome),
        "titulo_pap": _normalizar(titulo),
        "area_pap": area,
        "texto_extraido": texto_para_ia,
    }
